from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import io
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer
import faiss 
import numpy as np
import pandas as pd
import json
from queue import Queue
import threading
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import docx
import pptx
import time
import pickle
from typing import List, Tuple, Union, Dict
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from llama_index.llms.groq import Groq
import logging
from pinky import PlotAI
import matplotlib
matplotlib.use('Agg')
from fpdf import FPDF
import base64
import markdown
import openpyxl
from docx.shared import Inches
from fpdf import FPDF
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module="multiprocessing.resource_tracker")

# Set up logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)
from time import sleep
from datetime import datetime, timedelta
from threading import Lock

class RateLimiter:
    def __init__(self, requests_per_minute: int = 50):
        self.requests_per_minute = requests_per_minute
        self.requests = []
        self.lock = Lock()

    def wait_if_needed(self):
        """Wait if we've exceeded our rate limit"""
        now = datetime.now()
        minute_ago = now - timedelta(minutes=1)
        
        with self.lock:
            # Remove requests older than 1 minute
            self.requests = [req_time for req_time in self.requests if req_time > minute_ago]
            
            # If we've hit our limit, wait until we can make another request
            while len(self.requests) >= self.requests_per_minute:
                sleep(0.1)  # Wait 100ms before checking again
                now = datetime.now()
                minute_ago = now - timedelta(minutes=1)
                self.requests = [req_time for req_time in self.requests if req_time > minute_ago]
            
            # Add this request to our list
            self.requests.append(now)

class ReflectionPipeline:
    def __init__(self, groq_client):
        self.client = groq_client
        self.max_iterations = 4 # Reduced from 4 to 2
        self.improvement_threshold = 0.9
        self.rate_limiter = RateLimiter(requests_per_minute=45)  # Leave some buffer from the limit

    def make_api_call(self, messages: list, temperature: float = 0.7, max_tokens: int = 4096) -> str:
        """Make an API call with rate limiting"""
        try:
            self.rate_limiter.wait_if_needed()
            # Format the messages into a single prompt
            prompt = ""
            for msg in messages:
                if msg["role"] == "system":
                    prompt += f"System: {msg['content']}\n\n"
                else:
                    prompt += f"{msg['content']}\n\n"
            
            response = self.client.complete(
                prompt=prompt,
                temperature=temperature,
                max_tokens=max_tokens
            )
            return response.text
        except Exception as e:
            logger.error(f"API call failed: {str(e)}")
            raise

    def generate_combined_prompt(self, query: str, current_output: str) -> str:
        """Combine reflection and improvement into a single prompt"""
        return f"""Analyze and improve the following report. Focus on:
        1. Accuracy and factual correctness
        2. Clarity and organization
        3. Completeness of information
        4. Logical flow and coherence

        Original Query: {query}
        
        Current Report:
        {current_output}
        
        Provide an improved version of the report that addresses any issues you identify.
        Do not explain the changes, just provide the improved report:"""

    def assess_improvement_simple(self, original: str, improved: str) -> float:
        """
        Simple heuristic-based improvement assessment to avoid extra API call
        """
        # Length comparison (longer isn't always better, but too short usually indicates issues)
        len_ratio = len(improved) / max(len(original), 1)
        if len_ratio < 0.5 or len_ratio > 2:
            return 0.0

        # Structure improvements
        orig_sections = len([l for l in original.split('\n') if l.strip().startswith('#')])
        impr_sections = len([l for l in improved.split('\n') if l.strip().startswith('#')])
        structure_score = min(impr_sections / max(orig_sections, 1), 2) * 0.3

        # Content improvements (check for numerical data, lists, etc.)
        orig_numbers = len([w for w in original.split() if any(c.isdigit() for c in w)])
        impr_numbers = len([w for w in improved.split() if any(c.isdigit() for c in w)])
        numbers_score = min(impr_numbers / max(orig_numbers, 1), 2) * 0.3

        # Formatting improvements
        orig_formatting = len([l for l in original.split('\n') if l.strip().startswith(('-', '*', '1.'))])
        impr_formatting = len([l for l in improved.split('\n') if l.strip().startswith(('-', '*', '1.'))])
        formatting_score = min(impr_formatting / max(orig_formatting, 1), 2) * 0.4

        total_score = structure_score + numbers_score + formatting_score
        return min(max(total_score, 0), 1)

    def refine_output(self, query: str, initial_output: str) -> str:
        """
        Implement the reflection pipeline with reduced API calls and rate limiting
        """
        current_output = initial_output
        iteration = 0

        while iteration < self.max_iterations:
            try:
                # Combined reflection and improvement in single API call
                improved_output = self.make_api_call(
                    messages=[
                        {"role": "system", "content": "You are an expert report writer focused on accuracy and clarity."},
                        {"role": "user", "content": self.generate_combined_prompt(query, current_output)}
                    ]
                )

                # Use simple heuristic-based assessment
                improvement_score = self.assess_improvement_simple(current_output, improved_output)
                
                logger.info(f"Iteration {iteration + 1} improvement score: {improvement_score}")

                if improvement_score > self.improvement_threshold:
                    current_output = improved_output
                    iteration += 1
                    sleep(1)  # Add small delay between iterations
                else:
                    logger.info("Insufficient improvement, stopping refinement")
                    break

            except Exception as e:
                logger.error(f"Error in reflection pipeline iteration {iteration}: {str(e)}")
                break

        return current_output

class DocumentProcessor:
    def __init__(self):
        logger.info("Initializing DocumentProcessor")
        self.vector_dimension = 384
        
        # Initialize without GPU first
        self.use_gpu = False
        self.gpu_resource = None
        
        try:
            logger.info("Loading embedding model...")
            self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("Embedding model loaded successfully")
        except Exception as e:
            logger.error(f"Error loading embedding model: {str(e)}")
            raise

        # Create simple CPU index first
        try:
            # Start with a basic flat index for safety
            self.index = faiss.IndexFlatL2(self.vector_dimension)
            
            # Only attempt GPU initialization after basic index is created
            try:
                if faiss.get_num_gpus() > 0:
                    logger.info("GPU detected, attempting GPU initialization...")
                    self.gpu_resource = faiss.StandardGpuResources()
                    # Use conservative GPU memory settings
                    gpu_config = faiss.GpuIndexFlatConfig()
                    gpu_config.device = 0  # Use first GPU
                    gpu_config.useFloat16 = True  # Use float16 for memory efficiency
                    
                    # Safely convert to GPU index
                    try:
                        self.index = faiss.GpuIndexFlatL2(self.gpu_resource, self.vector_dimension, gpu_config)
                        self.use_gpu = True
                        logger.info("Successfully initialized GPU index")
                    except Exception as gpu_e:
                        logger.warning(f"GPU index initialization failed, falling back to CPU: {str(gpu_e)}")
                        # Recreate CPU index if GPU fails
                        self.index = faiss.IndexFlatL2(self.vector_dimension)
            except Exception as e:
                logger.warning(f"GPU detection failed, using CPU index: {str(e)}")
                
        except Exception as e:
            logger.error(f"Error creating FAISS index: {str(e)}")
            raise

        self.text_chunks = []
        self.chunk_sources = []
        self.processing_queue = Queue()
        self.processed_urls = set()

        # Initialize other components with memory safety in mind
        try:
            groq_api_key = os.getenv('GROQ_API_KEY')
            if not groq_api_key:
                raise ValueError("GROQ_API_KEY environment variable is not set")
            self.client = Groq(
                api_key=groq_api_key,
                model="llama-3.1-70b-versatile"  # Specify the model name
            )
            logger.info("Groq client initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing Groq client: {str(e)}")
            raise

        try:
            self.plot_ai = PlotAI(model_version="llama-3.1-70b-versatile", output_dir='static/plots')
            os.makedirs('static/plots', exist_ok=True)
            logger.info("PlotAI initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing PlotAI: {str(e)}")
            raise

        # Load existing index if available
        self.load_index()
        
        # Start processing thread
        self.start_processing_thread()
        self.rate_limiter = RateLimiter(requests_per_minute=45)

    def process_document(self, content: str, source_name: str) -> int:
        """Process a document with memory-safe batching."""
        try:
            logger.info(f"Processing document from source: {source_name}")
            chunks = self.split_text(content)
            if not chunks:
                logger.warning("No chunks created from document")
                return 0

            logger.info(f"Created {len(chunks)} chunks")
            
            # Use smaller batch size for memory safety
            batch_size = 100  # Reduced from 1000
            total_processed = 0
            
            for i in range(0, len(chunks), batch_size):
                batch_chunks = chunks[i:i + batch_size]
                try:
                    # Process in smaller batches
                    embeddings = self.embedding_model.encode(batch_chunks)
                    embeddings = embeddings.astype('float32')
                    faiss.normalize_L2(embeddings)
                    
                    # Add to index
                    self.index.add(embeddings)
                    self.text_chunks.extend(batch_chunks)
                    self.chunk_sources.extend([source_name] * len(batch_chunks))
                    total_processed += len(batch_chunks)
                    
                    # Periodic cleanup
                    if i % (batch_size * 10) == 0:
                        import gc
                        gc.collect()
                        
                except Exception as batch_e:
                    logger.error(f"Error processing batch {i//batch_size}: {str(batch_e)}")
                    continue

            self.save_index()
            logger.info(f"Successfully processed document, added {total_processed} chunks")
            return total_processed
            
        except Exception as e:
            logger.error(f"Error processing document {source_name}: {str(e)}")
            return 0

    def __del__(self):
        """Clean up GPU resources properly."""
        try:
            if self.use_gpu and self.gpu_resource is not None:
                del self.index
                del self.gpu_resource
                self.gpu_resource = None
                import gc
                gc.collect()
        except:
            pass

    def split_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:  # Increased overlap
        """Split text into overlapping chunks with increased overlap."""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap  # More overlap for better context preservation
            
        return chunks

    def extract_text_from_excel(self, file_content: bytes) -> str:
        """Extract text from Excel files."""
        try:
            workbook = openpyxl.load_workbook(io.BytesIO(file_content))
            text_content = []
            
            for sheet in workbook.sheetnames:
                worksheet = workbook[sheet]
                sheet_content = [f"Sheet: {sheet}"]
                
                for row in worksheet.iter_rows(values_only=True):
                    row_content = [str(cell) if cell is not None else '' for cell in row]
                    sheet_content.append(" | ".join(row_content))
                    
                text_content.append("\n".join(sheet_content))
                
            return "\n\n".join(text_content)
        except Exception as e:
            logger.error(f"Error extracting text from Excel file: {str(e)}")
            return ""

    def extract_text_from_csv(self, file_content: bytes) -> str:
        """Extract text from CSV files."""
        try:
            df = pd.read_csv(io.BytesIO(file_content))
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Error extracting text from CSV file: {str(e)}")
            return ""

    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text from various file formats."""
        try:
            extension = filename.lower().split('.')[-1]
            logger.info(f"Extracting text from file with extension: {extension}")
            
            if extension == 'pdf':
                with io.BytesIO(file_content) as pdf_file:
                    pdf_reader = PdfReader(pdf_file)
                    text = " ".join(page.extract_text() for page in pdf_reader.pages)
            elif extension in ['txt', 'json']:
                text = file_content.decode('utf-8')
                if extension == 'json':
                    json_data = json.loads(text)
                    text = json.dumps(json_data, indent=2)
            elif extension in ['docx', 'doc']:
                with io.BytesIO(file_content) as docx_file:
                    doc = docx.Document(docx_file)
                    text = " ".join(paragraph.text for paragraph in doc.paragraphs)
            elif extension == 'xlsx':
                text = self.extract_text_from_excel(file_content)
            elif extension == 'csv':
                text = self.extract_text_from_csv(file_content)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return ""
                
            logger.info(f"Successfully extracted text from {filename}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return ""

    def query_documents(self, query: str, k: int = 5) -> Tuple[str, List[Tuple[str, str]], pd.DataFrame, str]:
        """
        Query the document index with reduced strictness and return relevant answer, sources, structured data, and context.
        """
        try:
            logger.info(f"Processing query: {query}")
            query_embedding = self.embedding_model.encode([query])[0]
            query_embedding = query_embedding.reshape(1, -1).astype('float32')
            query_embedding = query_embedding + 1e-5
            faiss.normalize_L2(query_embedding)
            
            if self.index.ntotal == 0:
                logger.warning("No documents in index")
                return "No documents have been processed yet. Please add some documents first.", [], pd.DataFrame(), ""
            
            # Adjust search parameters for less strictness
            current_total = self.index.ntotal
            self.index.nprobe = min(60, max(1, current_total // 20))
            
            search_timeout = 30.0
            start_time = time.time()
            
            try:
                distances, indices = self.index.search(
                    query_embedding, 
                    min(k * 2, current_total)  # Retrieve more candidates
                )
            except Exception as e:
                logger.error(f"FAISS search failed: {str(e)}")
                return "Search operation failed. Please try again.", [], pd.DataFrame(), ""
            
            if time.time() - start_time > search_timeout:
                logger.warning("Search timeout exceeded")
                return "Search operation timed out. Please try a more specific query.", [], pd.DataFrame(), ""
                
            relevant_chunks = []
            structured_data = []
            
            # Process search results with more lenient threshold
            batch_size = 100
            for idx_batch in range(0, len(indices[0]), batch_size):
                batch_indices = indices[0][idx_batch:idx_batch + batch_size]
                
                for idx in batch_indices:
                    if idx != -1:
                        chunk_text = self.text_chunks[idx]
                        source = self.chunk_sources[idx]
                        
                        # Reduced relevance threshold from 0.3 to 0.2
                        score = 1 - distances[0][idx_batch + list(batch_indices).index(idx)]
                        
                        if score < 0.2:  # Lower threshold for more results
                            continue
                            
                        relevant_chunks.append({
                            'text': chunk_text,
                            'source': source,
                            'score': score
                        })
                        
                        try:
                            if chunk_text.strip().startswith('{'):
                                data_dict = json.loads(chunk_text)
                                structured_data.append(data_dict)
                            else:
                                import re
                                numbers = re.findall(r'(\w+):\s*(\d+(?:\.\d+)?)', chunk_text)
                                if numbers:
                                    data_dict = {label: float(value) for label, value in numbers}
                                    structured_data.append(data_dict)
                        except Exception as e:
                            logger.warning(f"Could not extract structured data from chunk: {str(e)}")
                            continue
            
            relevant_chunks.sort(key=lambda x: x['score'], reverse=True)
            
            if not relevant_chunks:
                logger.warning("No relevant chunks found")
                return "I couldn't find any relevant information to answer your question.", [], pd.DataFrame(), ""
            
            weighted_contexts = []
            total_weight = sum(chunk['score'] for chunk in relevant_chunks)
            
            for chunk in relevant_chunks:
                weight = chunk['score'] / total_weight
                weighted_context = f"""Content from {chunk['source']} (relevance: {chunk['score']:.2f}):
    {chunk['text']}"""
                weighted_contexts.append((weighted_context, weight))
            
            context = "\n\n".join(context for context, _ in weighted_contexts)
            
            prompt = f"""Based on the following context, please provide a well-structured report to answer the question.
            
            Context:
            {context}
            
            Question: {query}
            
            Report:"""
            
            try:
                self.rate_limiter.wait_if_needed()
                # Updated to use LlamaIndex's Groq interface
                system_prompt = """
                - You are a helpful report generating assistant called Zenit. 
                - Your task is to generate clear and well explained insightful reports based on parameters given in the query and the given context. 
                - The report must be professionally explained and should be worth enought to present before the authorities. Use of tables are prohibited. 
                - You may add your perspective and make the report in a professional and executive but in an explanative tone.
                """
                
                full_prompt = f"{system_prompt}\n\n{prompt}"
                
                initial_response = self.client.complete(
                    prompt=full_prompt,
                    temperature=0.7,
                    max_tokens=4096
                ).text
                
                # Initialize reflection pipeline with shared rate limiter
                reflection_pipeline = ReflectionPipeline(self.client)
                
                # Refine the report through reflection pipeline
                final_report = reflection_pipeline.refine_output(query, initial_response)
                
                df = pd.DataFrame(structured_data)
                sources = [(chunk['text'], chunk['source']) for chunk in relevant_chunks]
                
                logger.info("Successfully generated and refined report")
                return final_report, sources, df, final_report

            except Exception as e:
                logger.error(f"Error generating report: {str(e)}")
                return "Error generating report. Please try again.", [], pd.DataFrame(), ""
            
        except Exception as e:
            logger.error(f"Error in query_documents: {str(e)}")
            raise


    def load_index(self, directory: str = "./saved_index") -> bool:
        """Load the saved index if available."""
        try:
            if os.path.exists(directory):
                index_path = os.path.join(directory, "docs.index")
                chunks_path = os.path.join(directory, "chunks.pkl")
                
                if os.path.exists(index_path) and os.path.exists(chunks_path):
                    logger.info("Loading existing index...")
                    self.index = faiss.read_index(index_path)
                    with open(chunks_path, "rb") as f:
                        self.text_chunks, self.chunk_sources = pickle.load(f)
                    logger.info(f"Loaded index with {len(self.text_chunks)} chunks")
                    return True
            logger.info("No existing index found")
            return False
        except Exception as e:
            logger.error(f"Error loading index: {str(e)}")
            return False

    def save_index(self, directory: str = "./saved_index") -> None:
        """Save the current index."""
        try:
            os.makedirs(directory, exist_ok=True)
            logger.info(f"Saving index with {len(self.text_chunks)} chunks")
            faiss.write_index(self.index, os.path.join(directory, "docs.index"))
            with open(os.path.join(directory, "chunks.pkl"), "wb") as f:
                pickle.dump((self.text_chunks, self.chunk_sources), f)
            logger.info("Index saved successfully")
        except Exception as e:
            logger.error(f"Error saving index: {str(e)}")

    def start_processing_thread(self):
        """Start the background processing thread."""
        thread = threading.Thread(target=self.process_queue, daemon=True)
        thread.start()
        logger.info("Started processing thread")

    def process_queue(self):
        """Process documents in the queue."""
        while True:
            if not self.processing_queue.empty():
                content, source = self.processing_queue.get()
                self.process_document(content, source)
            time.sleep(1)

    def generate_pdf_report(self, report: str, plot_path: str = None) -> io.BytesIO:
        """Generate PDF report with optional plot using default fonts."""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Use built-in Helvetica font
            pdf.set_font('Helvetica', 'B', 16)
            pdf.cell(0, 10, "Generated Report", ln=True, align="C")
            pdf.ln(10)
            
            pdf.set_font('Helvetica', '', 12)
            
            # Add plot if available and if context has relevant data
            if plot_path and plot_path.startswith('/static/plots/'):
                try:
                    plot_abs_path = os.path.join(os.getcwd(), plot_path.lstrip('/'))
                    if os.path.exists(plot_abs_path):
                        img_width = 190
                        pdf.image(plot_abs_path, x=10, w=img_width)
                        pdf.ln(10)
                except Exception as e:
                    logger.error(f"Error adding plot to PDF: {str(e)}")
            
            for line in report.split('\n'):
                try:
                    if line.startswith('# '):
                        pdf.set_font('Helvetica', 'B', 14)
                        pdf.cell(0, 10, line.strip('# '), ln=True)
                        pdf.set_font('Helvetica', '', 12)
                    elif line.startswith('## '):
                        pdf.set_font('Helvetica', 'B', 13)
                        pdf.cell(0, 10, line.strip('# '), ln=True)
                        pdf.set_font('Helvetica', '', 12)
                    elif line.startswith('- '):
                        pdf.cell(10, 10, chr(149), ln=0)
                        pdf.multi_cell(0, 10, line.strip('- '))
                    elif line.strip():
                        cleaned_line = ''.join(char for char in line if ord(char) < 128)
                        pdf.multi_cell(0, 10, cleaned_line)
                except Exception as e:
                    logger.error(f"Error processing line in PDF: {str(e)}")
                    continue
            
            pdf_output = io.BytesIO()
            pdf.output(pdf_output)
            pdf_output.seek(0)
            
            return pdf_output
        except Exception as e:
            logger.error(f"Error generating PDF report: {str(e)}")
            raise

    def generate_html_report(self, report: str, plot_path: str = None) -> io.BytesIO:
        """Generate HTML report with optional plot using basic HTML/CSS."""
        try:
            html_content = markdown.markdown(report)
            
            css_styles = """
                body { 
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Arial, sans-serif; 
                    margin: 40px auto;
                    max-width: 800px;
                    line-height: 1.6;
                    color: #333;
                    padding: 0 20px;
                }
                h1 {
                    color: #2c3e50;
                    border-bottom: 2px solid #eee;
                    padding-bottom: 10px;
                }
                h2 {
                    color: #34495e;
                    margin-top: 30px;
                }
                img {
                    max-width: 100%;
                    height: auto;
                    margin: 20px 0;
                    border: 1px solid #eee;
                    border-radius: 5px;
                }
                ul, ol {
                    margin: 20px 0;
                    padding-left: 20px;
                }
                li {
                    margin: 10px 0;
                }
                p {
                    margin: 15px 0;
                }
                @media print {
                    body {
                        margin: 20px;
                    }
                    img {
                        max-width: 600px;
                    }
                }
            """
            
            plot_html = ""
            if plot_path and 'No relevant information' not in report:
                try:
                    plot_abs_path = os.path.join(os.getcwd(), plot_path.lstrip('/'))
                    if os.path.exists(plot_abs_path):
                        with open(plot_abs_path, 'rb') as img_file:
                            plot_base64 = base64.b64encode(img_file.read()).decode()
                            plot_html = f'<img src="data:image/png;base64,{plot_base64}" alt="Data Visualization">'
                except Exception as e:
                    logger.error(f"Error embedding plot in HTML: {str(e)}")
            
            html_template = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>Generated Report</title>
                <style>{css_styles}</style>
            </head>
            <body>
                <h1>Generated Report</h1>
                {plot_html}
                {html_content}
                <script>
                    const printButton = document.createElement('button');
                    printButton.innerHTML = 'Print Report';
                    printButton.style.cssText = 'position: fixed; bottom: 20px; right: 20px; padding: 10px 20px; background: #2c3e50; color: white; border: none; border-radius: 5px; cursor: pointer;';
                    printButton.onclick = () => window.print();
                    document.body.appendChild(printButton);
                </script>
            </body>
            </html>
            """
            
            html_output = io.BytesIO(html_template.encode('utf-8'))
            return html_output
        except Exception as e:
            logger.error(f"Error generating HTML report: {str(e)}")
            raise

    def generate_docx_report(self, report: str, plot_path: str = None) -> io.BytesIO:
        """Generate Word document report with optional plot."""
        try:
            doc = Document()
            
            doc.add_heading('Generated Report', 0)
            
            # Only add plot if context has relevant data
            if plot_path and plot_path.startswith('/static/plots/') and 'No relevant information' not in report:
                try:
                    plot_abs_path = os.path.join(os.getcwd(), plot_path.lstrip('/'))
                    if os.path.exists(plot_abs_path):
                        doc.add_picture(plot_abs_path, width=Inches(6))
                        doc.add_paragraph()
                except Exception as e:
                    logger.error(f"Error adding plot to Word document: {str(e)}")
            
            current_list = None
            for line in report.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line.strip('# '), 1)
                elif line.startswith('## '):
                    doc.add_heading(line.strip('# '), 2)
                elif line.startswith('- '):
                    if current_list is None or not current_list.style.name.startswith('List Bullet'):
                        current_list = doc.add_paragraph(style='List Bullet')
                    current_list.add_run(line.strip('- '))
                elif line.startswith('1. '):
                    if current_list is None or not current_list.style.name.startswith('List Number'):
                        current_list = doc.add_paragraph(style='List Number')
                    current_list.add_run(line[3:])
                elif line.strip():
                    current_list = None
                    doc.add_paragraph(line)
            
            docx_output = io.BytesIO()
            doc.save(docx_output)
            docx_output.seek(0)
            
            return docx_output
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            raise

    def generate_plot(self, context: str, query: str) -> Union[str, None]:
        """Generate a plot using PlotAI based on the context and query."""
        try:
            logger.info("Generating plot with PlotAI")
            
            # Check if context contains relevant numerical data
            if 'No relevant information' in context or not context.strip():
                logger.info("No relevant data found for plotting")
                return None
                
            # Look for numerical patterns in context
            import re
            numerical_patterns = re.findall(r'\d+(?:\.\d+)?', context)
            if len(numerical_patterns) < 4:  # Require minimum amount of numerical data
                logger.info("Insufficient numerical data for plotting")
                return None
            
            full_plot_query = f"""Generate a Bargraph for the query: {query}, dots alpha=0.3, the texts should not overlap, 
            give the dependant on Yaxis and the independant on Xaxis, the axises must be annoted correctly and also the 
            axises must be given the names correctly according to the asked query, strictly do not generate anything other 
            than the code for the graph, use textual annotations on each point according to the given context and the 
            annotations should not overlap: {context}"""
            
            logger.info(f"Full plot query: {full_plot_query}")
            
            plot_path = self.plot_ai.make(full_plot_query)
            
            if not os.path.exists(plot_path):
                logger.warning("Plot was not generated")
                return None
                
            logger.info(f"Plot generated at: {plot_path}")
            relative_path = os.path.relpath(plot_path, 'static/plots')
            
            return relative_path
            
        except Exception as e:
            logger.error(f"Error generating plot: {str(e)}")
            return None

# Initialize processor
try:
    processor = DocumentProcessor()
    logger.info("DocumentProcessor initialized successfully")
except Exception as e:
    logger.error(f"Error initializing DocumentProcessor: {str(e)}")
    raise

@app.route('/process', methods=['POST'])
def process_documents():
    try:
        logger.info("Received request to /process endpoint")
        
        if 'files' not in request.files and 'urls' not in request.form:
            logger.warning("No files or URLs in request")
            return jsonify({'error': 'No files or URLs provided'}), 400

        total_processed = 0

        if 'files' in request.files:
            files = request.files.getlist('files')
            logger.info(f"Processing {len(files)} files")
            
            for file in files:
                if file.filename:
                    try:
                        file_content = file.read()
                        text = processor.extract_text_from_file(file_content, file.filename)
                        if text:
                            processor.processing_queue.put((text, file.filename))
                            total_processed += 1
                    except Exception as e:
                        logger.error(f"Error processing file {file.filename}: {str(e)}")

        if 'urls' in request.form:
            urls = request.form['urls'].split('\n')
            logger.info(f"Processing {len(urls)} URLs")
            
            for url in urls:
                url = url.strip()
                if url:
                    try:
                        response = requests.get(url)
                        soup = BeautifulSoup(response.text, 'html.parser')
                        text = soup.get_text()
                        processor.processing_queue.put((text, url))
                        total_processed += 1
                    except Exception as e:
                        logger.error(f"Error processing URL {url}: {str(e)}")

        if total_processed > 0:
            return jsonify({
                'message': f'Successfully queued {total_processed} documents for processing'
            }), 200
        else:
            return jsonify({'error': 'No documents were processed successfully'}), 400

    except Exception as e:
        logger.error(f"Error in /process endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/query', methods=['POST'])
def query():
    try:
        logger.info("Received request to /query endpoint")
        data = request.json
        query = data.get('query')
        
        if not query:
            logger.warning("No query provided")
            return jsonify({'error': 'No query provided'}), 400

        logger.info(f"Processing query: {query}")
        report, sources, df, context = processor.query_documents(query)
        
        # Only generate plot if relevant data was found
        plot_path = None
        if 'No relevant information' not in report:
            try:
                plot_path = processor.generate_plot(context, query)
            except Exception as e:
                logger.warning(f"Error generating plot: {str(e)}")
        
        logger.info("Successfully generated report")
        return jsonify({
            'report': report,
            'sources': [{'text': text, 'source': source} for text, source in sources],
            'plot_path': f'/static/plots/{plot_path}' if plot_path else None
        }), 200

    except Exception as e:
        logger.error(f"Error in /query endpoint: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/report/<format>', methods=['POST'])
def generate_report(format):
    try:
        logger.info(f"Received request to generate {format} report")
        data = request.json
        report = data.get('report')
        plot_path = data.get('plotPath')
        
        if not report:
            logger.warning("No report provided")
            return jsonify({'error': 'No report provided'}), 400
            
        if format == 'html':
            output = processor.generate_html_report(report, plot_path)
            mimetype = 'text/html'
            filename = 'report.html'
        elif format == 'docx':
            output = processor.generate_docx_report(report, plot_path)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            filename = 'report.docx'
        else:
            return jsonify({'error': 'Unsupported format'}), 400
            
        return send_file(
            output,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"Error generating {format} report: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/static/plots/<path:filename>')
def serve_plot(filename):
    return send_file(f'static/plots/{filename}', mimetype='image/png')

if __name__ == '__main__':
    app.run(debug=True)